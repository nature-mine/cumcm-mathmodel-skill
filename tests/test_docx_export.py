from __future__ import annotations

import base64
import importlib.util
import json
import sys
import zipfile
from pathlib import Path
from types import ModuleType
from xml.etree import ElementTree

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "skills" / "cumcm-writing" / "scripts" / "docx_export.py"
ONE_PIXEL_PNG = (
    b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    b"+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _load_module(name: str = "docx_export") -> ModuleType:
    spec = importlib.util.spec_from_file_location(name, SCRIPT)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _prepare_workspace(tmp_path: Path) -> tuple[Path, str]:
    for dirname in ("code", "data", "figures", "input", "paper", "support"):
        (tmp_path / dirname).mkdir()
    source_code = (
        "from __future__ import annotations\n\n"
        "def solve(values: list[int]) -> int:\n"
        "\treturn sum(values)\n"
    )
    (tmp_path / "code" / "solver.py").write_text(source_code, encoding="utf-8")
    (tmp_path / "data" / "derived.csv").write_text("x,y\n1,2\n", encoding="utf-8")
    (tmp_path / "input" / "official.txt").write_text("official\n", encoding="utf-8")
    (tmp_path / "support" / "ai_usage.md").write_text(
        "AI 工具使用详情\n",
        encoding="utf-8",
    )
    (tmp_path / "figures" / "result.png").write_bytes(base64.b64decode(ONE_PIXEL_PNG))
    (tmp_path / "paper" / "paper.md").write_text(
        """\
# 测试论文

## 摘要

本文比较两个方案，冻结结果为 42。

关键词：确定性导出；支撑材料

## 1 问题重述

正文引用图 1 和表 1。

[[FIGURE id="FIG-A-01" path="figures/result.png" caption="结果曲线"]]

[[PLACEHOLDER id="FIG-G-01" lane="C" chapter="2.1" expected="反馈机制示意图"]]

[[TABLE caption="方案比较"]]
| 方案 | 得分 |
| --- | --- |
| A | 42 |
| B | 39 |
""",
        encoding="utf-8",
    )
    return tmp_path, source_code


def _paragraph_after(document: Document, heading: str) -> str:
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text == heading:
            return paragraphs[index + 1].text
    raise AssertionError(f"未找到段落：{heading}")


def test_export_preserves_source_and_classifies_support_materials(tmp_path: Path) -> None:
    workspace, source_code = _prepare_workspace(tmp_path)
    module = _load_module()

    summary = module.export_docx(
        workspace,
        "paper/paper.md",
        "paper/cumcm-paper.docx",
    )

    assert summary.source_code_files == 1
    assert summary.support_included == 3
    assert summary.support_excluded == 1
    assert (summary.figures, summary.placeholders, summary.tables) == (2, 1, 1)

    document = Document(workspace / summary.output)
    assert _paragraph_after(document, "源程序：code/solver.py") == source_code
    assert len(document.inline_shapes) == 1
    assert document.inline_shapes[0].width.mm <= 150
    assert any("FIG-G-01" in cell.text for table in document.tables for cell in table._cells)
    assert {"图 1　结果曲线", "图 2（占位）　反馈机制示意图", "表 1　方案比较"} <= {
        paragraph.text for paragraph in document.paragraphs
    }

    manifest = json.loads((workspace / summary.manifest).read_text(encoding="utf-8"))
    entries = {entry["path"]: entry for entry in manifest["entries"]}
    assert entries["input/official.txt"]["included"] is False
    assert entries["data/derived.csv"]["included"] is True
    assert entries["code/solver.py"]["included"] is True
    assert entries["support/ai_usage.md"]["included"] is True
    assert entries["support/ai_usage.md"]["category"] == "support"
    assert len(entries["code/solver.py"]["sha256"]) == 64


def test_export_applies_page_fonts_fields_and_is_byte_deterministic(tmp_path: Path) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_deterministic")

    first = module.export_docx(workspace, "paper/paper.md", "paper/first.docx")
    second = module.export_docx(workspace, "paper/paper.md", "paper/second.docx")
    first_path = workspace / first.output
    second_path = workspace / second.output
    assert first_path.read_bytes() == second_path.read_bytes()

    document = Document(first_path)
    section = document.sections[0]
    assert section.page_width.mm == pytest.approx(210, abs=0.1)
    assert section.page_height.mm == pytest.approx(297, abs=0.1)
    for margin in (
        section.top_margin,
        section.bottom_margin,
        section.left_margin,
        section.right_margin,
    ):
        assert margin.mm >= 25
    assert document.styles["Normal"].font.name == "Times New Roman"
    assert document.styles["Normal"].font.size.pt == 12

    with zipfile.ZipFile(first_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        footer_xml = archive.read("word/footer1.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        assert "<w:br w:type=\"page\"" in document_xml
        assert 'w:eastAsia="宋体"' in styles_xml
        assert " PAGE " in footer_xml
        assert 'w:updateFields w:val="true"' in settings_xml
        assert all(info.date_time == (2000, 1, 1, 0, 0, 0) for info in archive.infolist())


def test_export_writes_editable_equation_and_remains_byte_deterministic(
    tmp_path: Path,
) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_equation")
    source = workspace / "paper" / "equation.md"
    source.write_text(
        "# 公式导出\n\n"
        "## 摘要\n\n"
        "下式用于验证可编辑公式。\n\n"
        r'[[EQUATION latex="y_i=\frac{x^2+\sqrt{\alpha}}{2}"]]' + "\n",
        encoding="utf-8",
    )

    first = module.export_docx(workspace, "paper/equation.md", "paper/equation-first.docx")
    second = module.export_docx(workspace, "paper/equation.md", "paper/equation-second.docx")
    first_path = workspace / first.output
    second_path = workspace / second.output

    assert first_path.read_bytes() == second_path.read_bytes()
    with zipfile.ZipFile(first_path) as archive:
        document_bytes = archive.read("word/document.xml")
    document_xml = document_bytes.decode("utf-8")
    for token in ("<m:oMathPara>", "<m:f>", "<m:rad>", "<m:sSub>", "<m:sSup>"):
        assert token in document_xml
    assert "EQUATION latex" not in document_xml
    document_root = ElementTree.fromstring(document_bytes)
    namespaces = {
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    assert document_root.find(".//w:p/m:oMathPara", namespaces) is not None


def test_export_preserves_matrix_row_separator(tmp_path: Path) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_matrix_equation")
    source = workspace / "paper" / "matrix-equation.md"
    source.write_text(
        "# 矩阵公式\n\n"
        r'[[EQUATION latex="\begin{matrix}a\\b\end{matrix}"]]' "\n",
        encoding="utf-8",
    )

    summary = module.export_docx(
        workspace,
        "paper/matrix-equation.md",
        "paper/matrix-equation.docx",
    )

    with zipfile.ZipFile(workspace / summary.output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "<m:m>" in document_xml
    assert document_xml.count("<m:mr>") == 2


def test_parse_equation_directive_preserves_latex_backslashes() -> None:
    module = _load_module("docx_export_parse_equation")

    assert module._parse_directive(
        r'[[EQUATION latex="y=\frac{a}{b}"]]'
    ) == ("EQUATION", {"latex": r"y=\frac{a}{b}"})


def test_export_rejects_unquoted_equation_latex_without_writing_docx(
    tmp_path: Path,
) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_unquoted_equation")
    source = workspace / "paper" / "unquoted-equation.md"
    source.write_text(
        "# 无引号公式\n\n"
        r"[[EQUATION latex=\frac{1}{2}]]" "\n",
        encoding="utf-8",
    )

    with pytest.raises(
        ValueError,
        match="EQUATION 的 latex 字段必须用双引号包裹，值内不得含双引号字符",
    ):
        module.export_docx(
            workspace,
            "paper/unquoted-equation.md",
            "paper/unquoted-equation.docx",
        )
    assert not (workspace / "paper" / "unquoted-equation.docx").exists()


@pytest.mark.parametrize(
    "latex",
    (
        r"\frac{1}",
        r"\notacommand{x}",
        r"\sqrt{}",
        r"x_{1",
    ),
)
def test_export_rejects_invalid_latex_without_writing_docx(
    tmp_path: Path,
    latex: str,
) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module(f"docx_export_invalid_equation_{len(latex)}")
    source = workspace / "paper" / "invalid-equation.md"
    source.write_text(
        f'# 非法公式\n\n[[EQUATION latex="{latex}"]]\n',
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="EQUATION LaTeX 无效"):
        module.export_docx(
            workspace,
            "paper/invalid-equation.md",
            "paper/invalid-equation.docx",
        )
    assert not (workspace / "paper" / "invalid-equation.docx").exists()


def test_export_rejects_unsupported_mathml_without_writing_docx(tmp_path: Path) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_unsupported_equation")
    source = workspace / "paper" / "unsupported-equation.md"
    source.write_text(
        r'# 暂不支持的公式' "\n\n"
        r'[[EQUATION latex="\cancel{x}"]]' "\n",
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="暂不支持 MathML 元素 menclose"):
        module.export_docx(
            workspace,
            "paper/unsupported-equation.md",
            "paper/unsupported-equation.docx",
        )
    assert not (workspace / "paper" / "unsupported-equation.docx").exists()


def test_export_protects_negative_number_from_line_break(tmp_path: Path) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_negative_number")
    source = workspace / "paper" / "paper.md"
    source.write_text(
        "# 负数禁断行\n\n## 摘要\n\n"
        f"{'很长的摘要句子，' * 40}高钾扰动 ARI 的 5% 分位为 -0.0500，"
        "对应图号为 FIG-A-01。\n",
        encoding="utf-8",
    )

    summary = module.export_docx(workspace, "paper/paper.md", "paper/negative.docx")

    document = Document(workspace / summary.output)
    assert any("‑0.0500" in paragraph.text for paragraph in document.paragraphs)
    assert any("FIG-A-01" in paragraph.text for paragraph in document.paragraphs)
    with zipfile.ZipFile(workspace / summary.output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "‑0.0500" in document_xml
    assert "-0.0500" not in document_xml
    assert "FIG-A-01" in document_xml


def test_export_rejects_path_escape_and_invalid_directive(tmp_path: Path) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_invalid")

    with pytest.raises(ValueError, match="工作区相对路径"):
        module.export_docx(workspace, "../paper.md", "paper/out.docx")

    (workspace / "paper" / "bad.md").write_text(
        '[[PLACEHOLDER id="FIG-G-02" lane="A" chapter="2" expected="错误车道"]]\n',
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="lane 只能是 B 或 C"):
        module.export_docx(workspace, "paper/bad.md", "paper/bad.docx")
    assert not (workspace / "paper" / "bad.docx").exists()


def test_export_rejects_margin_below_verified_floor(tmp_path: Path) -> None:
    workspace, _ = _prepare_workspace(tmp_path)
    module = _load_module("docx_export_spec")
    bundled = json.loads(
        (ROOT / "skills/cumcm-writing/templates/cumcm-docx-spec.yaml").read_text(
            encoding="utf-8"
        )
    )
    bundled["page"]["margin_left_mm"] = 20
    (workspace / "unsafe-spec.yaml").write_text(
        json.dumps(bundled, ensure_ascii=False),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="页边距不得小于 25 mm"):
        module.export_docx(
            workspace,
            "paper/paper.md",
            "paper/unsafe.docx",
            spec_path="unsafe-spec.yaml",
        )

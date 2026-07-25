from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
EXPORTER = ROOT / "skills" / "cumcm-writing" / "scripts" / "docx_export.py"
TEMPLATE_SOURCE = (
    ROOT / "skills" / "cumcm-writing" / "templates" / "paper-template.md"
)
TEMPLATE_DOCX = (
    ROOT / "skills" / "cumcm-writing" / "templates" / "paper-template.docx"
)
PAPER_STRUCTURE = (
    ROOT / "skills" / "cumcm-writing" / "references" / "paper-structure.md"
)


def _prepare_template_workspace(path: Path) -> None:
    for dirname in ("paper", "input", "data", "code", "support"):
        (path / dirname).mkdir(parents=True)
    shutil.copyfile(TEMPLATE_SOURCE, path / "paper" / "paper.md")


def _run_export(
    workspace: Path,
    source: str,
    output: str,
) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [
            sys.executable,
            str(EXPORTER),
            "--workspace",
            str(workspace),
            "--source",
            source,
            "--output",
            output,
        ],
        cwd=ROOT,
        check=False,
        capture_output=True,
        text=True,
    )


def test_paper_template_regenerates_byte_for_byte_and_contains_contract(
    tmp_path: Path,
) -> None:
    template_text = TEMPLATE_SOURCE.read_text(encoding="utf-8")
    body_headings = re.findall(
        r"^### (\d+\.\s+.+)$",
        PAPER_STRUCTURE.read_text(encoding="utf-8"),
        re.MULTILINE,
    )
    template_headings = re.findall(r"^## (\d+\.\s+.+)$", template_text, re.MULTILINE)
    assert template_headings == body_headings
    assert len(body_headings) == 9
    assert "<!--" not in template_text
    assert len(re.findall(r"^\[\[TABLE\b", template_text, re.MULTILINE)) == 2
    assert len(re.findall(r'^\[\[PLACEHOLDER\b.*lane="B"', template_text, re.MULTILINE)) == 1
    assert len(re.findall(r'^\[\[PLACEHOLDER\b.*lane="C"', template_text, re.MULTILINE)) == 1
    assert re.search(r"^\[\[FIGURE\b", template_text, re.MULTILINE) is None
    assert "`[[FIGURE " in template_text
    for section in re.split(r"^## \d+\.\s+.+$", template_text, flags=re.MULTILINE)[1:]:
        placeholder_paragraphs = [
            line for line in section.splitlines() if line.startswith("【待替换】")
        ]
        assert 1 <= len(placeholder_paragraphs) <= 3

    _prepare_template_workspace(tmp_path)

    result = _run_export(tmp_path, "paper/paper.md", "paper/paper-template.docx")

    assert result.returncode == 0, result.stdout + result.stderr
    assert "LAYOUT figures=2 placeholders=2 tables=2" in result.stdout
    generated = tmp_path / "paper" / "paper-template.docx"
    assert generated.read_bytes() == TEMPLATE_DOCX.read_bytes()
    assert 10_000 <= generated.stat().st_size <= 200_000

    document = Document(generated)
    heading_texts = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1"
    ]
    assert heading_texts == ["摘要", *body_headings]
    assert heading_texts.count("9. 附录") == 1
    assert len(document.inline_shapes) == 0

    all_text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(cell.text for table in document.tables for cell in table._cells),
        ]
    )
    for token in (
        "问题：用最短篇幅界定总任务",
        "关键结果：填入来自 frozen 证据的量化结果",
        "关键词：【待替换】",
        "FIG-G-01",
        "车道：B",
        "FIG-G-02",
        "车道：C",
        '[[FIGURE id="FIG-A-01" path="figures/result.png"',
        "表 1　【待替换】模型假设及其依据",
        "表 2　【待替换】核心符号说明",
        "支撑材料文件清单",
        "完整源程序",
        "本论文没有用到程序。",
    ):
        assert token in all_text

    with zipfile.ZipFile(generated) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        footer_xml = archive.read("word/footer1.xml").decode("utf-8")
    assert '<w:br w:type="page"' in document_xml
    assert "<w:pageBreakBefore" in document_xml
    assert " PAGE " in footer_xml
    manifest = json.loads(
        (tmp_path / "paper" / "support_manifest.json").read_text(encoding="utf-8")
    )
    assert manifest["entries"] == []


def test_real_figure_directive_requires_an_existing_image(tmp_path: Path) -> None:
    paper = tmp_path / "paper"
    paper.mkdir()
    (paper / "paper.md").write_text(
        "# 图件失败样例\n\n"
        "## 摘要\n\n"
        "仅验证真实图指令。\n\n"
        '[[FIGURE id="FIG-A-01" path="figures/missing.png" caption="缺失图"]]\n',
        encoding="utf-8",
    )

    result = _run_export(tmp_path, "paper/paper.md", "paper/failed.docx")

    assert result.returncode == 2
    assert "ERROR" in result.stderr
    assert not (paper / "failed.docx").exists()

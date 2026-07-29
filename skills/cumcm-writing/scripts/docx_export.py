#!/usr/bin/env python3
"""Deterministically export a CUMCM Markdown paper to DOCX.

Adapted from:
  chinese-thesis-workbench-skill/scripts/docx/generate_thesis_docx.py

Copyright (c) 2026 Zyhsec

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.

This version narrows the Markdown contract for CUMCM, adds workspace path
safety, deterministic ZIP metadata, figure slots, automatic appendix source
import, editable OMML equations, and a support-material manifest. Runtime
dependencies python-docx and latex2mathml are supplied by the
cumcm-env-doctor Tier 0 contract.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import shlex
import sys
import zipfile
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from functools import lru_cache
from pathlib import Path
from types import SimpleNamespace
from typing import Any

DEFAULT_SPEC = Path(__file__).resolve().parents[1] / "templates" / "cumcm-docx-spec.yaml"
MATHML_TO_OMML_XSL = Path(__file__).with_name("mathml_to_omml.xsl")
IGNORED_PARTS = {
    ".git",
    ".mypy_cache",
    ".pytest_cache",
    ".ruff_cache",
    ".venv",
    "__pycache__",
}
GENERATED_SUFFIXES = {
    ".class",
    ".dll",
    ".dylib",
    ".exe",
    ".o",
    ".obj",
    ".pyc",
    ".pyo",
    ".so",
}
DIRECTIVE_PATTERN = re.compile(
    r"^\[\[(EQUATION|FIGURE|PLACEHOLDER|TABLE)\s+(.+)\]\]$"
)
LINK_PATTERN = re.compile(r"\[([^\]]+)]\([^)]+\)")
INLINE_MARK_PATTERN = re.compile(r"(\*\*|__|`)")
NEGATIVE_NUMBER_PATTERN = re.compile(r"(?<![A-Za-z0-9_])-(?=\d)")
FIXED_ZIP_TIME = (2000, 1, 1, 0, 0, 0)
MATHML_NAMESPACE = "http://www.w3.org/1998/Math/MathML"
SUPPORTED_MATHML_ELEMENTS = {
    "math",
    "mfrac",
    "mi",
    "mn",
    "mo",
    "mover",
    "mpadded",
    "mroot",
    "mrow",
    "mspace",
    "msqrt",
    "mstyle",
    "msub",
    "msubsup",
    "msup",
    "mtable",
    "mtd",
    "mtext",
    "mtr",
    "munder",
    "munderover",
}
MATHML_TOKEN_ELEMENTS = {"mi", "mn", "mo", "mtext"}
MATHML_FIXED_ARITY = {
    "mfrac": 2,
    "mover": 2,
    "mroot": 2,
    "msub": 2,
    "msubsup": 3,
    "msup": 2,
    "munder": 2,
    "munderover": 3,
}


@dataclass(frozen=True)
class SupportEntry:
    path: str
    category: str
    included: bool
    reason: str
    sha256: str


@dataclass(frozen=True)
class ExportSummary:
    output: str
    manifest: str
    source_code_files: int
    support_included: int
    support_excluded: int
    figures: int
    placeholders: int
    tables: int


def _load_docx_api() -> SimpleNamespace:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import (
            WD_CELL_VERTICAL_ALIGNMENT,
            WD_ROW_HEIGHT_RULE,
            WD_TABLE_ALIGNMENT,
        )
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Mm, Pt, RGBColor, Twips
    except ModuleNotFoundError as error:
        raise RuntimeError(
            "缺少 Tier 0 依赖 python-docx；先运行 cumcm-env-doctor 并按报告安装。"
        ) from error
    return SimpleNamespace(
        Document=Document,
        WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
        WD_BREAK=WD_BREAK,
        WD_CELL_VERTICAL_ALIGNMENT=WD_CELL_VERTICAL_ALIGNMENT,
        WD_ORIENT=WD_ORIENT,
        WD_ROW_HEIGHT_RULE=WD_ROW_HEIGHT_RULE,
        WD_TABLE_ALIGNMENT=WD_TABLE_ALIGNMENT,
        OxmlElement=OxmlElement,
        qn=qn,
        Cm=Cm,
        Mm=Mm,
        Pt=Pt,
        RGBColor=RGBColor,
        Twips=Twips,
    )


@lru_cache(maxsize=1)
def _load_equation_api() -> SimpleNamespace:
    try:
        from latex2mathml.converter import convert
        from lxml import etree
    except ModuleNotFoundError as error:
        raise RuntimeError(
            "缺少 Tier 0 公式依赖 latex2mathml；先运行 cumcm-env-doctor 并按报告安装。"
        ) from error

    if not MATHML_TO_OMML_XSL.is_file():
        raise RuntimeError(f"缺少公式转换资源：{MATHML_TO_OMML_XSL}")
    parser = etree.XMLParser(resolve_entities=False, no_network=True)
    try:
        stylesheet = etree.parse(str(MATHML_TO_OMML_XSL), parser)
        transform = etree.XSLT(stylesheet)
    except (OSError, etree.XMLSyntaxError, etree.XSLTParseError) as error:
        raise RuntimeError(f"公式转换资源无效：{MATHML_TO_OMML_XSL}") from error
    return SimpleNamespace(convert=convert, etree=etree, transform=transform)


def _safe_workspace_path(
    workspace: Path,
    raw_path: str | Path,
    *,
    must_exist: bool,
) -> Path:
    candidate = Path(raw_path)
    if candidate.is_absolute() or ".." in candidate.parts:
        raise ValueError(f"路径必须是工作区相对路径且不能包含 '..'：{raw_path}")

    root = workspace.resolve()
    resolved = (root / candidate).resolve()
    try:
        resolved.relative_to(root)
    except ValueError as error:
        raise ValueError(f"路径越出工作区：{raw_path}") from error
    if must_exist and not resolved.is_file():
        raise ValueError(f"输入文件不存在：{raw_path}")
    return resolved


def _load_spec(workspace: Path, raw_path: str | Path | None) -> dict[str, Any]:
    path = DEFAULT_SPEC if raw_path is None else _safe_workspace_path(
        workspace, raw_path, must_exist=True
    )
    try:
        spec = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, UnicodeError) as error:
        raise ValueError(f"DOCX 样式配置不是有效的 YAML 兼容 JSON：{path}") from error
    if not isinstance(spec, dict) or spec.get("schema_version") != 1:
        raise ValueError("DOCX 样式配置 schema_version 必须为 1")

    page = spec.get("page")
    if not isinstance(page, dict):
        raise ValueError("DOCX 样式配置缺少 page")
    margins = [page.get(f"margin_{side}_mm") for side in ("top", "bottom", "left", "right")]
    if any(not isinstance(value, (int, float)) or value < 25 for value in margins):
        raise ValueError("CUMCM 页边距不得小于 25 mm")
    return spec


def _is_scannable(path: Path, root: Path) -> bool:
    relative = path.relative_to(root)
    if any(part.startswith(".") or part in IGNORED_PARTS for part in relative.parts):
        return False
    return path.suffix.lower() not in GENERATED_SUFFIXES


def _scan_directory(workspace: Path, dirname: str) -> list[Path]:
    root = workspace.resolve()
    directory = (root / dirname).resolve()
    if not directory.is_dir():
        return []
    try:
        directory.relative_to(root)
    except ValueError as error:
        raise ValueError(f"{dirname}/ 越出工作区") from error

    files: list[Path] = []
    for path in directory.rglob("*"):
        if not path.is_file() or not _is_scannable(path, root):
            continue
        resolved = path.resolve()
        try:
            resolved.relative_to(root)
        except ValueError as error:
            raise ValueError(f"扫描到越出工作区的链接：{path.relative_to(root)}") from error
        files.append(resolved)
    return sorted(files, key=lambda item: item.relative_to(root).as_posix())


def _file_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def build_support_manifest(workspace: Path) -> list[SupportEntry]:
    workspace = workspace.resolve()
    entries: list[SupportEntry] = []
    categories = (
        ("input", False, "赛题原始附件，不进入支撑材料压缩包"),
        ("data", True, "自主查阅或派生数据，进入支撑材料压缩包"),
        ("code", True, "完整源程序，进入支撑材料压缩包并自动导入附录"),
        ("support", True, "竞赛要求的说明性支撑材料，进入支撑材料压缩包"),
    )
    for category, included, reason in categories:
        for path in _scan_directory(workspace, category):
            entries.append(
                SupportEntry(
                    path=path.relative_to(workspace).as_posix(),
                    category=category,
                    included=included,
                    reason=reason,
                    sha256=_file_sha256(path),
                )
            )
    return entries


def _read_source_files(workspace: Path) -> list[tuple[str, str]]:
    source_files: list[tuple[str, str]] = []
    for path in _scan_directory(workspace, "code"):
        relative = path.relative_to(workspace).as_posix()
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeError as error:
            raise ValueError(f"源程序必须是 UTF-8 文本：{relative}") from error
        if "\x00" in text:
            raise ValueError(f"源程序不能含 NUL 字符：{relative}")
        source_files.append((relative, text))
    return source_files


def _write_manifest(path: Path, entries: list[SupportEntry]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "schema_version": 1,
        "entries": [asdict(entry) for entry in entries],
    }
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _set_rfonts(element: Any, api: SimpleNamespace, east_asia: str, latin: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = api.OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(api.qn("w:eastAsia"), east_asia)
    r_fonts.set(api.qn("w:ascii"), latin)
    r_fonts.set(api.qn("w:hAnsi"), latin)


def _format_run(
    run: Any,
    api: SimpleNamespace,
    *,
    east_asia: str,
    latin: str,
    size_pt: float,
    bold: bool | None = None,
) -> None:
    run.font.name = latin
    run.font.size = api.Pt(size_pt)
    if bold is not None:
        run.bold = bold
    _set_rfonts(run._element, api, east_asia, latin)


def _configure_styles(document: Any, api: SimpleNamespace, spec: dict[str, Any]) -> None:
    fonts = spec["fonts"]
    sizes = spec["sizes_pt"]
    paragraph = spec["paragraph"]
    definitions = {
        "Normal": (fonts["body_east_asia"], fonts["body_latin"], sizes["body"], False),
        "Title": (fonts["heading_east_asia"], fonts["body_latin"], sizes["title"], True),
        "Heading 1": (
            fonts["heading_east_asia"],
            fonts["body_latin"],
            sizes["heading_1"],
            True,
        ),
        "Heading 2": (
            fonts["subheading_east_asia"],
            fonts["body_latin"],
            sizes["heading_2"],
            True,
        ),
        "Heading 3": (
            fonts["subheading_east_asia"],
            fonts["body_latin"],
            sizes["heading_3"],
            True,
        ),
        "Caption": (
            fonts["body_east_asia"],
            fonts["body_latin"],
            sizes["caption"],
            False,
        ),
    }
    for name, (east_asia, latin, size, bold) in definitions.items():
        style = document.styles[name]
        style.font.name = latin
        style.font.size = api.Pt(size)
        style.font.bold = bold
        style.font.color.rgb = api.RGBColor(0, 0, 0)
        _set_rfonts(style.element, api, east_asia, latin)

    title_style = document.styles["Title"]
    title_style.paragraph_format.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    title_p_pr = title_style.element.get_or_add_pPr()
    title_border = title_p_pr.find(api.qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    normal_format = document.styles["Normal"].paragraph_format
    normal_format.line_spacing = paragraph["line_spacing"]
    normal_format.first_line_indent = api.Cm(paragraph["first_line_indent_cm"])
    normal_format.space_after = api.Pt(paragraph["space_after_pt"])
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        heading_format = document.styles[name].paragraph_format
        heading_format.keep_with_next = True
        heading_format.space_before = api.Pt(paragraph["heading_space_before_pt"])
        heading_format.space_after = api.Pt(paragraph["heading_space_after_pt"])
    document.styles["Heading 1"].paragraph_format.alignment = api.WD_ALIGN_PARAGRAPH.CENTER


def _configure_document(document: Any, api: SimpleNamespace, spec: dict[str, Any]) -> None:
    page = spec["page"]
    section = document.sections[0]
    section.orientation = api.WD_ORIENT.PORTRAIT
    section.page_width = api.Mm(page["width_mm"])
    section.page_height = api.Mm(page["height_mm"])
    section.top_margin = _minimum_mm(api, page["margin_top_mm"])
    section.bottom_margin = _minimum_mm(api, page["margin_bottom_mm"])
    section.left_margin = _minimum_mm(api, page["margin_left_mm"])
    section.right_margin = _minimum_mm(api, page["margin_right_mm"])

    fixed_time = datetime(2000, 1, 1, tzinfo=UTC)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""


def _minimum_mm(api: SimpleNamespace, value: float) -> Any:
    """Round a minimum dimension upward to Word's integral twip grid."""

    return api.Twips(math.ceil(value * 1440 / 25.4))


def _add_update_fields_setting(document: Any, api: SimpleNamespace) -> None:
    settings = document.settings.element
    existing = settings.find(api.qn("w:updateFields"))
    if existing is None:
        existing = api.OxmlElement("w:updateFields")
        _insert_before_successor(
            settings,
            existing,
            api,
            (
                "w:hdrShapeDefaults",
                "w:footnotePr",
                "w:endnotePr",
                "w:compat",
                "w:docVars",
                "w:rsids",
                "m:mathPr",
                "w:themeFontLang",
                "w:clrSchemeMapping",
                "w:doNotAutoCompressPictures",
            ),
        )
    existing.set(api.qn("w:val"), "true")


def _insert_before_successor(
    parent: Any,
    element: Any,
    api: SimpleNamespace,
    successor_tags: tuple[str, ...],
) -> None:
    successors = {api.qn(tag) for tag in successor_tags}
    for child in parent:
        if child.tag in successors:
            child.addprevious(element)
            return
    parent.append(element)


def _add_page_number_footer(document: Any, api: SimpleNamespace, spec: dict[str, Any]) -> None:
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    field_parts = (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:t", {}, "1"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    )
    for tag, attributes, text in field_parts:
        run = paragraph.add_run()
        _format_run(
            run,
            api,
            east_asia=spec["fonts"]["body_east_asia"],
            latin=spec["fonts"]["body_latin"],
            size_pt=spec["sizes_pt"]["footer"],
        )
        element = api.OxmlElement(tag)
        for key, value in attributes.items():
            element.set(api.qn(key), value)
        if text is not None:
            element.text = text
        run._r.append(element)


def _clean_inline(text: str) -> str:
    cleaned = INLINE_MARK_PATTERN.sub("", LINK_PATTERN.sub(r"\1", text))
    return NEGATIVE_NUMBER_PATTERN.sub("\N{NON-BREAKING HYPHEN}", cleaned).strip()


def _parse_directive(line: str) -> tuple[str, dict[str, str]] | None:
    match = DIRECTIVE_PATTERN.fullmatch(line.strip())
    if match is None:
        return None
    kind, raw_fields = match.groups()
    fields: dict[str, str] = {}
    for token in shlex.split(raw_fields):
        key, separator, value = token.partition("=")
        if not separator or not key or not value or key in fields:
            raise ValueError(f"DOCX 指令字段无效：{line}")
        fields[key] = value
    return kind, fields


def _require_fields(kind: str, fields: dict[str, str], required: set[str]) -> None:
    missing = sorted(required - fields.keys())
    extra = sorted(fields.keys() - required)
    if missing or extra:
        raise ValueError(f"{kind} 指令字段不匹配：missing={missing}, extra={extra}")


def _equation_error(reason: str) -> ValueError:
    return ValueError(f"EQUATION LaTeX 无效：{reason}")


def _validate_latex_source(latex: str) -> str:
    value = latex.strip()
    if not value:
        raise _equation_error("表达式为空")
    if "$" in value:
        raise _equation_error("latex 字段不得包含 $ 定界符")
    if any(ord(character) < 32 for character in value):
        raise _equation_error("表达式含控制字符")

    depth = 0
    escaped = False
    for character in value:
        if escaped:
            escaped = False
            continue
        if character == "\\":
            escaped = True
        elif character == "{":
            depth += 1
        elif character == "}":
            depth -= 1
            if depth < 0:
                raise _equation_error("花括号不配对")
    if depth:
        raise _equation_error("花括号不配对")
    return value


def _mathml_local_name(element: Any, etree: Any) -> str:
    return str(etree.QName(element).localname)


def _mathml_has_content(element: Any, etree: Any) -> bool:
    return any(
        _mathml_local_name(descendant, etree) in MATHML_TOKEN_ELEMENTS
        and bool((descendant.text or "").strip())
        for descendant in element.iter()
    )


def _validate_mathml(root: Any, etree: Any) -> None:
    elements = [element for element in root.iter() if isinstance(element.tag, str)]
    for element in elements:
        qualified_name = etree.QName(element)
        if qualified_name.namespace != MATHML_NAMESPACE:
            raise _equation_error("转换结果含非 MathML 元素")
        local_name = str(qualified_name.localname)
        if local_name not in SUPPORTED_MATHML_ELEMENTS:
            raise _equation_error(f"暂不支持 MathML 元素 {local_name}")
        if element.text and "\\" in element.text:
            raise _equation_error(f"未知或不支持的 LaTeX 命令 {element.text}")

    if _mathml_local_name(root, etree) != "math" or not _mathml_has_content(root, etree):
        raise _equation_error("表达式没有可转换内容")

    for element in elements:
        local_name = _mathml_local_name(element, etree)
        children = [child for child in element if isinstance(child.tag, str)]
        required_arity = MATHML_FIXED_ARITY.get(local_name)
        if required_arity is not None and len(children) != required_arity:
            raise _equation_error(
                f"{local_name} 需要 {required_arity} 个操作数，实际为 {len(children)}"
            )
        if required_arity is not None and any(
            not _mathml_has_content(child, etree) for child in children
        ):
            raise _equation_error(f"{local_name} 含空操作数")
        if local_name == "msqrt" and (
            not children or not all(_mathml_has_content(child, etree) for child in children)
        ):
            raise _equation_error("msqrt 含空操作数")
        if local_name == "mtable":
            invalid_rows = (
                not children
                or any(_mathml_local_name(child, etree) != "mtr" for child in children)
            )
            if invalid_rows:
                raise _equation_error("矩阵缺少有效行")
            column_counts = {
                len([cell for cell in row if isinstance(cell.tag, str)])
                for row in children
            }
            if len(column_counts) != 1:
                raise _equation_error("矩阵各行列数不一致")
        if local_name == "mtr" and (
            not children or any(_mathml_local_name(child, etree) != "mtd" for child in children)
        ):
            raise _equation_error("矩阵行缺少有效单元格")
        if local_name == "mtd" and not _mathml_has_content(element, etree):
            raise _equation_error("矩阵含空单元格")


def _add_equation(
    document: Any,
    api: SimpleNamespace,
    fields: dict[str, str],
) -> None:
    _require_fields("EQUATION", fields, {"latex"})
    latex = _validate_latex_source(fields["latex"])
    equation_api = _load_equation_api()
    try:
        mathml = equation_api.convert(latex, display="block")
    except Exception as error:
        raise _equation_error(f"转换失败（{type(error).__name__}）") from error

    parser = equation_api.etree.XMLParser(resolve_entities=False, no_network=True)
    try:
        root = equation_api.etree.fromstring(mathml.encode("utf-8"), parser)
    except equation_api.etree.XMLSyntaxError as error:
        raise _equation_error("转换结果不是有效 MathML") from error
    _validate_mathml(root, equation_api.etree)

    try:
        transformed = equation_api.transform(root)
    except equation_api.etree.XSLTApplyError as error:
        raise _equation_error("MathML 无法转换为可编辑 OMML") from error
    omml = transformed.getroot()
    if omml is None:
        raise _equation_error("MathML 转换未产生 OMML")

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph._p.append(omml)


def _set_picture_alt(inline_shape: Any, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)


def _add_figure(
    document: Any,
    api: SimpleNamespace,
    workspace: Path,
    fields: dict[str, str],
    spec: dict[str, Any],
    number: int,
) -> None:
    _require_fields("FIGURE", fields, {"id", "path", "caption"})
    image_path = _safe_workspace_path(workspace, fields["path"], must_exist=True)
    if image_path.suffix.lower() not in {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff"}:
        raise ValueError(f"FIGURE 不支持图像格式：{fields['path']}")

    page = spec["page"]
    body_width = page["width_mm"] - page["margin_left_mm"] - page["margin_right_mm"]
    width = min(body_width, spec["figures"]["max_width_mm"])
    paragraph = document.add_paragraph()
    paragraph.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    shape = paragraph.add_run().add_picture(str(image_path), width=api.Mm(width))
    _set_picture_alt(shape, f"{fields['id']}：{fields['caption']}")

    caption = document.add_paragraph(style="Caption")
    caption.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = None
    caption.add_run(f"图 {number}　{fields['caption']}")


def _shade_cell(cell: Any, api: SimpleNamespace, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(api.qn("w:shd"))
    if shading is None:
        shading = api.OxmlElement("w:shd")
        _insert_before_successor(
            tc_pr,
            shading,
            api,
            (
                "w:noWrap",
                "w:tcMar",
                "w:textDirection",
                "w:tcFitText",
                "w:vAlign",
                "w:hideMark",
                "w:headers",
            ),
        )
    shading.set(api.qn("w:val"), "clear")
    shading.set(api.qn("w:color"), "auto")
    shading.set(api.qn("w:fill"), color)


def _set_table_width(table: Any, api: SimpleNamespace, width_mm: float) -> None:
    preferred_width = table._tbl.tblPr.find(api.qn("w:tblW"))
    if preferred_width is None:
        preferred_width = api.OxmlElement("w:tblW")
        _insert_before_successor(
            table._tbl.tblPr,
            preferred_width,
            api,
            (
                "w:jc",
                "w:tblCellSpacing",
                "w:tblInd",
                "w:tblBorders",
                "w:shd",
                "w:tblLayout",
                "w:tblCellMar",
                "w:tblLook",
            ),
        )
    preferred_width.set(api.qn("w:type"), "dxa")
    preferred_width.set(api.qn("w:w"), str(round(width_mm * 1440 / 25.4)))


def _add_placeholder(
    document: Any,
    api: SimpleNamespace,
    fields: dict[str, str],
    spec: dict[str, Any],
    number: int,
) -> None:
    _require_fields("PLACEHOLDER", fields, {"id", "lane", "chapter", "expected"})
    if fields["lane"] not in {"B", "C"}:
        raise ValueError("PLACEHOLDER lane 只能是 B 或 C")
    table = document.add_table(rows=1, cols=1)
    table.alignment = api.WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_width(table, api, spec["figures"]["max_width_mm"])
    cell = table.cell(0, 0)
    cell.width = api.Mm(spec["figures"]["max_width_mm"])
    cell.vertical_alignment = api.WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.rows[0].height = api.Mm(spec["figures"]["placeholder_height_mm"])
    table.rows[0].height_rule = api.WD_ROW_HEIGHT_RULE.AT_LEAST
    _shade_cell(cell, api, spec["figures"]["placeholder_fill"])

    paragraph = cell.paragraphs[0]
    paragraph.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.add_run(
        f"图示占位槽\nfigure_id：{fields['id']}\n车道：{fields['lane']}\n"
        f"所在章节：{fields['chapter']}\n期望内容：{fields['expected']}"
    )
    caption = document.add_paragraph(style="Caption")
    caption.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = None
    caption.add_run(f"图 {number}（占位）　{fields['expected']}")


def _set_table_borders(table: Any, api: SimpleNamespace) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(api.qn("w:tblBorders"))
    if borders is None:
        borders = api.OxmlElement("w:tblBorders")
        _insert_before_successor(
            tbl_pr,
            borders,
            api,
            (
                "w:shd",
                "w:tblLayout",
                "w:tblCellMar",
                "w:tblLook",
                "w:tblCaption",
                "w:tblDescription",
            ),
        )
    for edge, value, size in (
        ("top", "single", "8"),
        ("left", "nil", "0"),
        ("bottom", "single", "8"),
        ("right", "nil", "0"),
        ("insideH", "nil", "0"),
        ("insideV", "nil", "0"),
    ):
        element = api.OxmlElement(f"w:{edge}")
        element.set(api.qn("w:val"), value)
        element.set(api.qn("w:sz"), size)
        element.set(api.qn("w:space"), "0")
        element.set(api.qn("w:color"), "000000")
        borders.append(element)


def _set_cell_bottom_border(cell: Any, api: SimpleNamespace) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(api.qn("w:tcBorders"))
    if borders is None:
        borders = api.OxmlElement("w:tcBorders")
        _insert_before_successor(
            tc_pr,
            borders,
            api,
            (
                "w:shd",
                "w:noWrap",
                "w:tcMar",
                "w:textDirection",
                "w:tcFitText",
                "w:vAlign",
                "w:hideMark",
                "w:headers",
            ),
        )
    bottom = api.OxmlElement("w:bottom")
    bottom.set(api.qn("w:val"), "single")
    bottom.set(api.qn("w:sz"), "4")
    bottom.set(api.qn("w:space"), "0")
    bottom.set(api.qn("w:color"), "000000")
    borders.append(bottom)


def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [_clean_inline(cell) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    if len(rows) < 2 or len({len(row) for row in rows}) != 1:
        raise ValueError("Markdown 表格至少需要表头和一行数据，且每行列数必须一致")
    return rows


def _add_table(
    document: Any,
    api: SimpleNamespace,
    lines: list[str],
    caption_text: str,
    number: int,
) -> None:
    rows = _parse_markdown_table(lines)
    caption = document.add_paragraph(style="Caption")
    caption.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = None
    caption.add_run(f"表 {number}　{caption_text}")

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = api.WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, api)
    for row_index, row in enumerate(rows):
        for cell_index, text in enumerate(row):
            cell = table.cell(row_index, cell_index)
            cell.text = text
            cell.vertical_alignment = api.WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = None
                if row_index == 0:
                    _set_cell_bottom_border(cell, api)
                    for run in paragraph.runs:
                        run.bold = True


def _add_code_block(document: Any, api: SimpleNamespace, text: str, spec: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = api.Pt(2)
    paragraph.paragraph_format.keep_together = False
    run = paragraph.add_run(text)
    _format_run(
        run,
        api,
        east_asia=spec["fonts"]["code_east_asia"],
        latin=spec["fonts"]["code_latin"],
        size_pt=spec["sizes_pt"]["code"],
    )


def _add_manifest_appendix(
    document: Any,
    api: SimpleNamespace,
    entries: list[SupportEntry],
    source_files: list[tuple[str, str]],
    spec: dict[str, Any],
    *,
    appendix_heading_present: bool = False,
) -> None:
    if not appendix_heading_present:
        document.add_page_break()
        document.add_heading("附录", level=1)
    document.add_heading("支撑材料文件清单", level=2)
    if entries:
        table = document.add_table(rows=1, cols=4)
        table.alignment = api.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_borders(table, api)
        column_widths_mm = (45, 25, 25, 65)
        _set_table_width(table, api, sum(column_widths_mm))
        for index, width in enumerate(column_widths_mm):
            table.columns[index].width = api.Mm(width)
        headers = ("路径", "类别", "进入压缩包", "说明")
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            cell.width = api.Mm(column_widths_mm[index])
        for entry in entries:
            cells = table.add_row().cells
            values = (
                entry.path,
                entry.category,
                "是" if entry.included else "否",
                entry.reason,
            )
            for index, value in enumerate(values):
                cells[index].text = value
                cells[index].width = api.Mm(column_widths_mm[index])
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = api.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = None
                    if row_index == 0:
                        _set_cell_bottom_border(cell, api)
                        for run in paragraph.runs:
                            run.bold = True
    else:
        document.add_paragraph("未扫描到 input/、data/ 或 code/ 文件。")

    document.add_heading("完整源程序", level=2)
    if not source_files:
        document.add_paragraph("本论文没有用到程序。")
        return
    for relative, text in source_files:
        heading = document.add_heading(f"源程序：{relative}", level=3)
        heading.paragraph_format.page_break_before = True
        _add_code_block(document, api, text, spec)


def _render_markdown(
    document: Any,
    api: SimpleNamespace,
    workspace: Path,
    source: Path,
    spec: dict[str, Any],
) -> tuple[int, int, int, bool]:
    lines = source.read_text(encoding="utf-8").splitlines()
    figure_count = 0
    placeholder_count = 0
    table_count = 0
    pending_table_caption: str | None = None
    abstract_seen = False
    body_started = False
    appendix_heading_seen = False
    current_section = ""
    in_code = False
    code_lines: list[str] = []
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()
        if in_code:
            if stripped.startswith("```"):
                _add_code_block(document, api, "\n".join(code_lines), spec)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw_line)
            index += 1
            continue
        if stripped.startswith("```"):
            in_code = True
            index += 1
            continue
        if not stripped or stripped == "---":
            index += 1
            continue

        directive = _parse_directive(stripped)
        if directive is not None:
            kind, fields = directive
            if kind == "EQUATION":
                _add_equation(document, api, fields)
            elif kind == "FIGURE":
                figure_count += 1
                _add_figure(document, api, workspace, fields, spec, figure_count)
            elif kind == "PLACEHOLDER":
                figure_count += 1
                placeholder_count += 1
                _add_placeholder(document, api, fields, spec, figure_count)
            else:
                _require_fields("TABLE", fields, {"caption"})
                if pending_table_caption is not None:
                    raise ValueError("连续 TABLE 指令之间缺少 Markdown 表格")
                pending_table_caption = fields["caption"]
            index += 1
            continue

        if stripped.startswith("|"):
            if pending_table_caption is None:
                raise ValueError("Markdown 表格前必须有 [[TABLE caption=\"...\"]] 指令")
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            table_count += 1
            _add_table(document, api, table_lines, pending_table_caption, table_count)
            pending_table_caption = None
            continue
        if pending_table_caption is not None:
            raise ValueError("TABLE 指令后必须紧跟 Markdown 表格")

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading is not None:
            marks, text = heading.groups()
            text = _clean_inline(text)
            if len(marks) == 1:
                paragraph = document.add_paragraph(text, style="Title")
                paragraph.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = None
            else:
                normalized = text.replace(" ", "").lower()
                semantic_heading = re.sub(
                    r"^(?:\d+|[一二三四五六七八九十]+)[.、．]?",
                    "",
                    normalized,
                )
                if normalized == "摘要":
                    abstract_seen = True
                elif len(marks) == 2 and abstract_seen and not body_started:
                    body_started = True
                    paragraph = document.add_paragraph()
                    paragraph.add_run().add_break(api.WD_BREAK.PAGE)
                level = min(len(marks) - 1, 3)
                paragraph = document.add_heading(text, level=level)
                if normalized in {"摘要", "参考文献", "附录"}:
                    paragraph.alignment = api.WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = None
                if len(marks) == 2 and semantic_heading == "附录":
                    appendix_heading_seen = True
                    paragraph.paragraph_format.page_break_before = True
                if len(marks) == 2:
                    current_section = normalized
            index += 1
            continue

        paragraph = document.add_paragraph(_clean_inline(stripped))
        if stripped.startswith(("- ", "* ")):
            paragraph.text = _clean_inline(stripped[2:])
            paragraph.style = "List Bullet"
            paragraph.paragraph_format.first_line_indent = None
        elif re.match(r"^\d+[.)]\s+", stripped):
            paragraph.text = _clean_inline(re.sub(r"^\d+[.)]\s+", "", stripped))
            paragraph.style = "List Number"
            paragraph.paragraph_format.first_line_indent = None
        elif current_section in {"摘要", "abstract"} or stripped.startswith(
            ("关键词：", "关键词:", "Keywords:", "Keywords：")
        ):
            paragraph.paragraph_format.first_line_indent = None
        index += 1

    if in_code:
        raise ValueError("Markdown 代码围栏未闭合")
    if pending_table_caption is not None:
        raise ValueError("TABLE 指令后缺少 Markdown 表格")
    return figure_count, placeholder_count, table_count, appendix_heading_seen


def _normalize_docx_zip(path: Path) -> None:
    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as target:
        for name in sorted(source.namelist()):
            original = source.getinfo(name)
            info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            target.writestr(info, source.read(name))
    temporary.replace(path)


def export_docx(
    workspace: Path,
    source_path: str | Path,
    output_path: str | Path,
    *,
    manifest_path: str | Path | None = None,
    spec_path: str | Path | None = None,
) -> ExportSummary:
    workspace = workspace.resolve()
    if not workspace.is_dir():
        raise ValueError(f"工作区不存在：{workspace}")
    source = _safe_workspace_path(workspace, source_path, must_exist=True)
    output = _safe_workspace_path(workspace, output_path, must_exist=False)
    if output.suffix.lower() != ".docx":
        raise ValueError("输出文件扩展名必须是 .docx")
    manifest = _safe_workspace_path(
        workspace,
        manifest_path or output.with_name("support_manifest.json").relative_to(workspace),
        must_exist=False,
    )
    if manifest.suffix.lower() != ".json":
        raise ValueError("支撑材料清单扩展名必须是 .json")

    api = _load_docx_api()
    spec = _load_spec(workspace, spec_path)
    entries = build_support_manifest(workspace)
    source_files = _read_source_files(workspace)
    document = api.Document()
    _configure_document(document, api, spec)
    _configure_styles(document, api, spec)
    _add_update_fields_setting(document, api)
    _add_page_number_footer(document, api, spec)
    figures, placeholders, tables, appendix_heading_seen = _render_markdown(
        document, api, workspace, source, spec
    )
    _add_manifest_appendix(
        document,
        api,
        entries,
        source_files,
        spec,
        appendix_heading_present=appendix_heading_seen,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    _normalize_docx_zip(output)
    _write_manifest(manifest, entries)

    included = sum(entry.included for entry in entries)
    return ExportSummary(
        output=output.relative_to(workspace).as_posix(),
        manifest=manifest.relative_to(workspace).as_posix(),
        source_code_files=len(source_files),
        support_included=included,
        support_excluded=len(entries) - included,
        figures=figures,
        placeholders=placeholders,
        tables=tables,
    )


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="确定性导出 CUMCM DOCX、附录源程序和支撑材料清单。")
    parser.add_argument("--workspace", type=Path, required=True, help="竞赛工作区根目录")
    parser.add_argument("--source", required=True, help="工作区相对 Markdown 路径")
    parser.add_argument("--output", default="paper/cumcm-paper.docx", help="工作区相对 DOCX 路径")
    parser.add_argument("--manifest", default=None, help="工作区相对清单 JSON 路径")
    parser.add_argument("--spec", default=None, help="工作区相对样式覆盖文件")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    try:
        summary = export_docx(
            args.workspace,
            args.source,
            args.output,
            manifest_path=args.manifest,
            spec_path=args.spec,
        )
    except (OSError, RuntimeError, UnicodeError, ValueError) as error:
        print(f"ERROR：{error}", file=sys.stderr)
        return 2

    print(f"WROTE {summary.output}")
    print(f"MANIFEST {summary.manifest}")
    print(f"SOURCE_CODE_FILES {summary.source_code_files}")
    print(
        "SUPPORT_MATERIALS "
        f"included={summary.support_included} excluded={summary.support_excluded}"
    )
    print(
        f"LAYOUT figures={summary.figures} "
        f"placeholders={summary.placeholders} tables={summary.tables}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
